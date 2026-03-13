import os
from docx import Document  # Changed from 'from docx import Document'
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor

# Output Word document name
OUTPUT_DOC = 'loom_code.docx'

def add_code_paragraph(document, code):
    """
    Adds a formatted code block to the Word document.
    """
    # Add a paragraph for the code
    code_paragraph = document.add_paragraph()
    
    # Set the style to 'Code' if it exists, otherwise create it
    styles = document.styles
    if 'Code' not in styles:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = 'Courier New'
        code_font.size = Pt(10)
        code_font.color.rgb = RGBColor(0, 0, 0)
    else:
        code_style = styles['Code']
    
    code_paragraph.style = 'Code'
    
    # Add the code text
    run = code_paragraph.add_run(code)
    
    # Optionally, add shading to highlight the code block
    p = code_paragraph._element
    pPr = p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'D3D3D3')  # Light gray background
    pPr.append(shd)

def generate_document():
    """
    Generates the Word document with all .go files in the current directory.
    """
    # Create a new Word document
    document = Document()
    
    # Define document title
    document.add_heading('LOOM AI FRAMEWORK Go Code Documentation', 0)
    
    # Get all .go files in the current directory
    go_files = [f for f in os.listdir('.') if f.endswith('.go')]
    
    if not go_files:
        print("No .go files found in the current directory.")
        return
    
    # Sort the files alphabetically for consistent ordering
    go_files.sort()
    
    for i, go_file in enumerate(go_files):
        # Add a page break before each new file (except the first)
        if i > 0:
            document.add_page_break()
        
        # Add the file name as a heading
        document.add_heading(go_file, level=1)
        
        # Read the .go file content
        try:
            with open(go_file, 'r', encoding='utf-8') as f:
                code = f.read()
            
            # Add the code block to the document
            add_code_paragraph(document, code)
        except Exception as e:
            print(f"Error reading {go_file}: {str(e)}")
            continue
    
    # Save the document, replacing if it already exists
    document.save(OUTPUT_DOC)
    print(f"Documentation successfully generated and saved as {OUTPUT_DOC} with {len(go_files)} Go files")

if __name__ == "__main__":
    generate_document()
