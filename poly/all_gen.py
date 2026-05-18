import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Output Word document name (written next to this script)
OUTPUT_DOC = "loom_code_poly.docx"

# Source kinds included in the documentation bundle
SOURCE_EXTENSIONS = {".go", ".s"}

# Directories to skip when walking the tree
SKIP_DIR_NAMES = {".git", "__pycache__", "vendor", ".cache", "node_modules"}

ROOT = Path(__file__).resolve().parent


def add_code_paragraph(document, code):
    """Adds a formatted code block to the Word document."""
    code_paragraph = document.add_paragraph()

    styles = document.styles
    if "Code" not in styles:
        code_style = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = "Courier New"
        code_font.size = Pt(10)
        code_font.color.rgb = RGBColor(0, 0, 0)

    code_paragraph.style = "Code"
    code_paragraph.add_run(code)

    p = code_paragraph._element
    p_pr = p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "D3D3D3")
    p_pr.append(shd)


def collect_source_files(root: Path) -> list[Path]:
    """Return sorted paths (relative to root) for all .go and .s files under root."""
    found: list[Path] = []
    for dirpath, dirnames, filenames in os.walk(root):
        dirnames[:] = [d for d in dirnames if d not in SKIP_DIR_NAMES and not d.startswith(".")]
        for name in filenames:
            path = Path(dirpath) / name
            if path.suffix in SOURCE_EXTENSIONS:
                found.append(path.relative_to(root))
    return sorted(found, key=lambda p: p.as_posix())


def section_key(rel: Path) -> str:
    """Group files for section headings: root poly vs asm/<subpackage>."""
    parts = rel.parts
    if not parts or parts[0] != "asm":
        return "poly"
    if len(parts) >= 2:
        return f"asm/{parts[1]}"
    return "asm"


def section_title(key: str) -> str:
    if key == "poly":
        return "Core poly (package root)"
    if key.startswith("asm/"):
        return f"Assembly — {key.replace('/', os.sep)}"
    return key


def generate_document():
    """Generate a Word document with all .go and .s sources under loom/poly."""
    document = Document()
    document.add_heading("LOOM poly — Go & assembly source documentation", 0)

    sources = collect_source_files(ROOT)
    if not sources:
        print(f"No .go or .s files found under {ROOT}")
        return

    go_count = sum(1 for p in sources if p.suffix == ".go")
    asm_count = sum(1 for p in sources if p.suffix == ".s")
    document.add_paragraph(
        f"Aggregates {len(sources)} source files from the poly tree "
        f"({go_count} Go, {asm_count} assembly .s), including asm/dot, asm/dense, and asm/matmul."
    )

    prev_section = None
    for i, rel in enumerate(sources):
        rel_str = rel.as_posix()
        sec = section_key(rel)

        if sec != prev_section:
            document.add_heading(section_title(sec), level=1)
            prev_section = sec
        elif i > 0:
            document.add_page_break()

        document.add_heading(rel_str, level=2)

        try:
            code = (ROOT / rel).read_text(encoding="utf-8")
            add_code_paragraph(document, code)
        except OSError as exc:
            print(f"Error reading {rel_str}: {exc}")
            continue

    out_path = ROOT / OUTPUT_DOC
    document.save(out_path)
    print(
        f"Documentation saved as {out_path} "
        f"({len(sources)} files: {go_count} .go, {asm_count} .s)"
    )


if __name__ == "__main__":
    generate_document()
