#!/usr/bin/env python3
"""
Build a single master Word document from all Markdown files in this directory.

Usage (from repo root or docs/):
    pip install python-docx
    python docs/make_docx.py
    python docs/make_docx.py -o docs/loom_master.docx

Output default: docs/master.docx
"""

from __future__ import annotations

import argparse
import os
import re
import sys
from datetime import datetime, timezone

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_OUTPUT = "master.docx"

# Read order: index first, then same sequence as docs/index.md (remaining files appended A–Z).
MD_ORDER = [
    "index.md",
    "overview.md",
    "deployment.md",
    "donate_compute.md",
    "tanhi.md",
    "numerical_types.md",
    "layers.md",
    "dispatch.md",
    "training.md",
    "gpu.md",
    "step.md",
    "dna.md",
    "evolution.md",
    "softmax.md",
    "serialization.md",
    "parallel_sequential.md",
    "quantization.md",
    "transformer.md",
    "quick_reference.md",
    "testing_and_validation.md",
    "bitnet_cpu.md",
]

SKIP_FILES = {"make_docx.py", "requirements-docx.txt"}


def sorted_md_files(directory: str) -> list[str]:
    all_md = {
        f
        for f in os.listdir(directory)
        if f.endswith(".md") and f not in SKIP_FILES
    }
    ordered = [f for f in MD_ORDER if f in all_md]
    rest = sorted(all_md - set(ordered))
    return ordered + rest


def ensure_code_style(doc: Document):
    styles = doc.styles
    if "Code Block" in styles:
        return styles["Code Block"]
    style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = "Courier New"
    font.size = Pt(9)
    font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    pf = style.paragraph_format
    pf.left_indent = Inches(0.3)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    style.element.pPr.append(shd)
    return style


def add_code_block(doc: Document, code_style, text: str):
    for line in text.split("\n"):
        p = doc.add_paragraph(line, style=code_style)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)


def add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def parse_inline(run_text: str):
    segments: list[tuple[str, bool, bool, bool]] = []
    pattern = re.compile(r"`([^`]+)`")
    last = 0
    for m in pattern.finditer(run_text):
        if m.start() > last:
            segments.append((run_text[last : m.start()], False, False, False))
        segments.append((m.group(1), False, False, True))
        last = m.end()
    if last < len(run_text):
        segments.append((run_text[last:], False, False, False))

    result: list[tuple[str, bool, bool, bool]] = []
    for text, _b, _i, c in segments:
        if c:
            result.append((text, False, False, True))
            continue
        parts = re.split(r"\*\*(.+?)\*\*", text)
        for idx, part in enumerate(parts):
            if idx % 2 == 1:
                result.append((part, True, False, False))
            elif part:
                result.append((part, False, False, False))
    return result


def add_paragraph_with_inline(doc: Document, text: str, style_name: str = "Normal"):
    p = doc.add_paragraph(style=style_name)
    for segment, bold, italic, code in parse_inline(text):
        if not segment:
            continue
        run = p.add_run(segment)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4F)
    return p


def parse_and_add_table(doc: Document, lines: list[str]):
    rows = []
    for line in lines:
        line = line.strip()
        if re.match(r"^\|[-:| ]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= col_count:
                break
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            for segment, bold, italic, code in parse_inline(cell_text):
                if not segment:
                    continue
                run = p.add_run(segment)
                run.bold = bold or (r_idx == 0)
                if code:
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
    doc.add_paragraph()


def process_md(doc: Document, filepath: str, code_style):
    with open(filepath, encoding="utf-8") as f:
        lines = f.readlines()

    in_code = False
    code_lines: list[str] = []
    in_table = False
    table_lines: list[str] = []
    i = 0

    while i < len(lines):
        line = lines[i].rstrip("\n")

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, code_style, "\n".join(code_lines))
                in_code = False
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("|"):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        if in_table:
            parse_and_add_table(doc, table_lines)
            in_table = False
            table_lines = []

        stripped = line.strip()

        if re.match(r"^---+$", stripped):
            add_horizontal_rule(doc)
            i += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("> ") or stripped.startswith(">!"):
            text = re.sub(r"^>!?\s*\[!NOTE\]\s*", "", stripped)
            text = re.sub(r"^>\s*", "", text)
            p = add_paragraph_with_inline(doc, text)
            p.paragraph_format.left_indent = Inches(0.3)
            if p.runs:
                p.runs[0].italic = True
            i += 1
            continue

        list_match = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)", line)
        if list_match:
            indent = len(list_match.group(1))
            text = list_match.group(3)
            style = "List Bullet" if indent == 0 else "List Bullet 2"
            add_paragraph_with_inline(doc, text, style)
            i += 1
            continue

        add_paragraph_with_inline(doc, stripped)
        i += 1

    if in_table:
        parse_and_add_table(doc, table_lines)


def build_master_docx(output_path: str) -> int:
    md_files = sorted_md_files(SCRIPT_DIR)
    if not md_files:
        print("No .md files found.", file=sys.stderr)
        return 1

    doc = Document()
    doc.add_heading("Loom / poly — Master Documentation", 0)
    meta = doc.add_paragraph()
    meta.add_run(
        f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')} "
        f"from {len(md_files)} markdown files in docs/"
    ).italic = True

    code_style = ensure_code_style(doc)

    for idx, filename in enumerate(md_files):
        filepath = os.path.join(SCRIPT_DIR, filename)
        print(f"  {filename}")
        if idx > 0:
            doc.add_page_break()
        doc.add_heading(filename, level=1)
        process_md(doc, filepath, code_style)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    print(f"\nSaved: {output_path}  ({len(md_files)} chapters)")
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(description="Merge docs/*.md into one Word document.")
    parser.add_argument(
        "-o",
        "--output",
        default=os.path.join(SCRIPT_DIR, DEFAULT_OUTPUT),
        help=f"Output .docx path (default: docs/{DEFAULT_OUTPUT})",
    )
    args = parser.parse_args()
    output = os.path.abspath(args.output)

    try:
        from docx import Document  # noqa: F401 — import check
    except ImportError:
        print("Missing dependency: pip install python-docx", file=sys.stderr)
        return 1

    return build_master_docx(output)


if __name__ == "__main__":
    raise SystemExit(main())
